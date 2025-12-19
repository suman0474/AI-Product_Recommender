"""
File Extraction Utilities
Extracts text from various file formats: PDF, DOCX, DOC, TXT, CSV, Images (OCR)
"""

import logging
import io
import csv
from typing import Optional, Dict, Any
from PIL import Image
import fitz  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX

# Initialize logger first
logger = logging.getLogger(__name__)

# Make EasyOCR optional (only needed for image OCR)
try:
    import easyocr
    import os
    
    # Initialize EasyOCR reader
    # Using English by default, can add more languages as needed: ['en', 'es', 'fr', etc.]
    # gpu=False for CPU-only environments (Railway/most cloud deployments)
    # download_enabled=True allows downloading models on first use
    logger.info("Initializing EasyOCR reader...")
    
    try:
        # Initialize with English language support
        # Set gpu=False for CPU-only environments
        # The model will be downloaded on first use (~100MB for English)
        easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
        EASYOCR_AVAILABLE = True
        logger.info("✓ EasyOCR initialized successfully with English language support")
    except Exception as e:
        EASYOCR_AVAILABLE = False
        logger.warning(f"⚠ Failed to initialize EasyOCR reader: {e}")
        easyocr_reader = None
        
except ImportError:
    EASYOCR_AVAILABLE = False
    easyocr_reader = None
    logger.warning("⚠ easyocr not installed - image OCR will not be available")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_bytes: PDF file as bytes
        
    Returns:
        Extracted text
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        doc.close()
        
        extracted_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from PDF ({len(text_parts)} pages)")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file_bytes: DOCX file as bytes
        
    Returns:
        Extracted text
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        extracted_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        return ""


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from image using OCR (EasyOCR)
    
    Args:
        file_bytes: Image file as bytes
        
    Returns:
        Extracted text via OCR
    """
    if not EASYOCR_AVAILABLE or easyocr_reader is None:
        error_msg = (
            "EasyOCR is not installed or failed to initialize. "
            "To enable image text extraction:\n"
            "1. Install easyocr: pip install easyocr\n"
            "2. Restart the backend server\n"
            "Note: EasyOCR will download language models (~100MB for English) on first use."
        )
        logger.error(error_msg)
        raise RuntimeError(error_msg)
    
    try:
        # Open and validate image
        image = Image.open(io.BytesIO(file_bytes))
        
        # Convert to RGB if necessary (EasyOCR works best with RGB)
        if image.mode not in ('RGB', 'L'):
            logger.info(f"Converting image from {image.mode} to RGB for OCR")
            image = image.convert('RGB')
        
        # Perform OCR with EasyOCR
        # readtext returns a list of tuples: (bbox, text, confidence)
        # We extract just the text from each detection
        logger.info("Performing OCR with EasyOCR...")
        results = easyocr_reader.readtext(file_bytes)
        
        # Extract text from results
        # Each result is a tuple: (bbox, text, confidence)
        extracted_texts = [text for (bbox, text, confidence) in results]
        extracted_text = '\n'.join(extracted_texts).strip()
        
        logger.info(f"Extracted {len(extracted_text)} characters from image via EasyOCR ({len(results)} text regions detected)")
        
        if not extracted_text:
            logger.warning("OCR completed but no text was found in the image")
            raise RuntimeError("No text found in image. The image may be blank, too low quality, or contain only graphics.")
        
        return extracted_text
        
    except RuntimeError:
        # Re-raise RuntimeError as-is (already formatted)
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from image with EasyOCR: {e}")
        logger.exception(e)
        raise RuntimeError(f"Image OCR failed: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extract text from plain text file
    
    Args:
        file_bytes: Text file as bytes
        
    Returns:
        Decoded text
    """
    try:
        # Try UTF-8 first
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            text = file_bytes.decode('latin-1')
        
        logger.info(f"Extracted {len(text)} characters from text file")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Failed to extract text from text file: {e}")
        return ""


def extract_text_from_csv(file_bytes: bytes) -> str:
    """
    Extract text from CSV file
    
    Args:
        file_bytes: CSV file as bytes
        
    Returns:
        Text representation of CSV data
    """
    try:
        # Try UTF-8 first, then fallback
        try:
            text_content = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = file_bytes.decode('latin-1')
            except UnicodeDecodeError:
                text_content = file_bytes.decode('cp1252')
        
        # Parse CSV
        csv_reader = csv.reader(io.StringIO(text_content))
        rows = list(csv_reader)
        
        if not rows:
            logger.warning("CSV file is empty")
            return ""
        
        # Extract headers
        headers = rows[0] if rows else []
        
        # Convert to readable text format for LLM processing
        text_parts = []
        text_parts.append(f"CSV File with {len(rows)} rows and {len(headers)} columns")
        text_parts.append(f"Headers: {', '.join(headers)}")
        text_parts.append("")
        
        # Add each row as a structured entry
        for i, row in enumerate(rows[1:], start=1):  # Skip header row
            row_data = []
            for j, value in enumerate(row):
                if j < len(headers):
                    header = headers[j]
                    if value.strip():  # Only include non-empty values
                        row_data.append(f"{header}: {value.strip()}")
            
            if row_data:
                text_parts.append(f"Row {i}: {' | '.join(row_data)}")
        
        extracted_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from CSV ({len(rows)} rows)")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from CSV file: {e}")
        return ""


def extract_text_from_excel(file_bytes: bytes) -> str:
    """
    Extract text from Excel file (xlsx, xls)
    
    Args:
        file_bytes: Excel file as bytes
        
    Returns:
        Text representation of Excel data
    """
    try:
        import openpyxl
        
        workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        text_parts = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
                
            headers = [str(cell) if cell else "" for cell in rows[0]]
            text_parts.append(f"Headers: {', '.join(headers)}")
            
            for i, row in enumerate(rows[1:], start=1):
                row_data = []
                for j, cell in enumerate(row):
                    if cell is not None and j < len(headers):
                        header = headers[j]
                        row_data.append(f"{header}: {cell}")
                
                if row_data:
                    text_parts.append(f"Row {i}: {' | '.join(row_data)}")
            
            text_parts.append("")
        
        extracted_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(extracted_text)} characters from Excel")
        return extracted_text
        
    except ImportError:
        logger.error("openpyxl not installed - Excel extraction not available")
        return ""
    except Exception as e:
        logger.error(f"Failed to extract text from Excel file: {e}")
        return ""


def extract_text_from_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract text from uploaded file based on file type
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename with extension
        
    Returns:
        Dict with extracted text and metadata
    """
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    logger.info(f"Extracting text from file: {filename} (type: {file_ext})")
    
    extracted_text = ""
    file_type = ""
    
    # PDF files
    if file_ext == 'pdf':
        file_type = "PDF"
        extracted_text = extract_text_from_pdf(file_bytes)
    
    # DOCX files
    elif file_ext == 'docx':
        file_type = "DOCX"
        extracted_text = extract_text_from_docx(file_bytes)
    
    # DOC files (older Word format - requires conversion or different library)
    elif file_ext == 'doc':
        file_type = "DOC"
        logger.warning("DOC format not fully supported. Please convert to DOCX or PDF.")
        extracted_text = ""
    
    # Text files
    elif file_ext in ['txt', 'text', 'log']:
        file_type = "TEXT"
        extracted_text = extract_text_from_txt(file_bytes)
    
    # CSV files
    elif file_ext == 'csv':
        file_type = "CSV"
        extracted_text = extract_text_from_csv(file_bytes)
    
    # Excel files
    elif file_ext in ['xlsx', 'xls']:
        file_type = "EXCEL"
        extracted_text = extract_text_from_excel(file_bytes)
    
    # Image files (OCR)
    elif file_ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'tif']:
        file_type = "IMAGE"
        try:
            extracted_text = extract_text_from_image(file_bytes)
        except RuntimeError as e:
            logger.warning(f"Image OCR failed for {filename}: {e}")
            return {
                "filename": filename,
                "file_type": file_type,
                "extracted_text": "",
                "character_count": 0,
                "success": False,
                "error": str(e)
            }
    
    else:
        logger.warning(f"Unsupported file type: {file_ext}")
        file_type = "UNSUPPORTED"
    
    return {
        "filename": filename,
        "file_type": file_type,
        "extracted_text": extracted_text,
        "character_count": len(extracted_text),
        "success": len(extracted_text) > 0
    }
