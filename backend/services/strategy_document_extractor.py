"""
Strategy Document Extractor Service
====================================
Extracts vendor strategy data from uploaded documents using LLM.

Supports: CSV, Excel (XLS/XLSX), PDF, DOCX, TXT

Extracts 4 fields per vendor:
- vendor_name: Name of the vendor
- category: Product category (e.g., "Flow Instruments")
- subcategory: Product subcategory (e.g., "Ultrasonic Flow Meters")
- strategy: Procurement strategy description
"""

import logging
import io
import json
from typing import List, Dict, Any, Optional
import pandas as pd

logger = logging.getLogger(__name__)

# Try to import document parsing libraries
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")

try:
    # Try pypdf first (newer), then fall back to PyPDF2 (legacy)
    try:
        from pypdf import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PdfReader = None
    logger.warning("PDF library not available. Install with: pip install pypdf")


class StrategyDocumentExtractor:
    """
    Extracts structured vendor strategy data from documents using LLM.
    """

    def __init__(self):
        """Initialize the extractor with LLM support."""
        self.llm = None
        self._initialize_llm()

    def _initialize_llm(self):
        """Initialize LLM for extraction."""
        try:
            from services.llm.fallback import create_llm_with_fallback
            self.llm = create_llm_with_fallback(
                model="gemini-2.5-flash",
                temperature=0.0,  # Deterministic for extraction
                timeout=36000
            )
            logger.info("[STRATEGY_EXTRACTOR] LLM initialized successfully")
        except Exception as e:
            logger.error(f"[STRATEGY_EXTRACTOR] Failed to initialize LLM: {e}")
            self.llm = None

    def extract_from_file(
        self,
        file_bytes: bytes,
        filename: str,
        content_type: str,
        user_id: int = None
    ) -> List[Dict[str, str]]:
        """
        Extract vendor strategy data from a file.

        Args:
            file_bytes: File content as bytes
            filename: Original filename
            content_type: MIME type
            user_id: User ID for user-scoped keyword standardization

        Returns:
            List of dicts with keys: vendor_name, category, subcategory, strategy
            (with standardized fields added)

        Raises:
            ValueError: If file format is not supported or parsing fails
        """
        logger.info(f"[STRATEGY_EXTRACTOR] Extracting from {filename} ({content_type})")

        # Determine file type and parse
        text_content = None

        if filename.endswith('.csv') or content_type == 'text/csv':
            text_content = self._parse_csv(file_bytes)
        elif filename.endswith(('.xlsx', '.xls')) or 'spreadsheet' in content_type:
            text_content = self._parse_excel(file_bytes)
        elif filename.endswith('.pdf') or content_type == 'application/pdf':
            text_content = self._parse_pdf(file_bytes)
        elif filename.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text_content = self._parse_docx(file_bytes)
        elif filename.endswith('.txt') or content_type == 'text/plain':
            text_content = file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {filename}")

        if not text_content or len(text_content.strip()) < 10:
            raise ValueError("Document appears to be empty or too short")

        logger.info(f"[STRATEGY_EXTRACTOR] Parsed document: {len(text_content)} characters")

        # Extract structured data using LLM
        vendor_data = self._extract_with_llm(text_content, filename, user_id=user_id)

        logger.info(f"[STRATEGY_EXTRACTOR] Extracted {len(vendor_data)} vendor records")
        return vendor_data

    def _parse_csv(self, file_bytes: bytes) -> str:
        """Parse CSV file to text representation."""
        try:
            # Try UTF-8 first
            df = pd.read_csv(io.BytesIO(file_bytes), encoding='utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            df = pd.read_csv(io.BytesIO(file_bytes), encoding='latin-1')

        # Convert to text format preserving structure
        return df.to_string(index=False)

    def _parse_excel(self, file_bytes: bytes) -> str:
        """Parse Excel file to text representation."""
        df = pd.read_excel(io.BytesIO(file_bytes), engine='openpyxl')
        return df.to_string(index=False)

    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Parse PDF file to text."""
        if not PDF_AVAILABLE or not PdfReader:
            raise ValueError("PDF parsing not available. Install pypdf.")

        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for page in pdf_reader.pages:
            text_parts.append(page.extract_text())

        return '\n'.join(text_parts)

    def _parse_docx(self, file_bytes: bytes) -> str:
        """Parse DOCX file to text."""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX parsing not available. Install python-docx.")

        doc = DocxDocument(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                text_parts.append(row_text)

        return '\n'.join(text_parts)

    def _extract_with_llm(
        self,
        document_text: str,
        filename: str,
        user_id: int = None
    ) -> List[Dict[str, str]]:
        """
        Use LLM to extract structured vendor data from document text and standardize keywords.
        Supports processing large documents in chunks.

        Args:
            document_text: The parsed document text
            filename: Original filename for context
            user_id: User ID for user-scoped keyword standardization

        Returns:
            List of vendor records with 4 fields each (plus standardized fields)
        """
        if not self.llm:
            logger.error("[STRATEGY_EXTRACTOR] LLM not available, cannot extract")
            raise RuntimeError("LLM not initialized. Cannot extract vendor data.")

        # Chunking configuration (Gemini Flash processes very fast, large chunks are fine)
        chunk_size = 300000  # 300k chars per chunk
        overlap = 2000       # 2k chars overlap

        # Create chunks
        chunks = []
        if len(document_text) <= chunk_size:
            chunks.append(document_text)
        else:
            logger.info(f"[STRATEGY_EXTRACTOR] Document large ({len(document_text)} chars), splitting into chunks...")
            start = 0
            while start < len(document_text):
                end = min(start + chunk_size, len(document_text))
                chunks.append(document_text[start:end])
                if end == len(document_text):
                    break
                start = end - overlap  # Move forward with overlap

        all_vendor_data = []

        try:
            for i, chunk in enumerate(chunks):
                logger.info(f"[STRATEGY_EXTRACTOR] Processing chunk {i + 1}/{len(chunks)} ({len(chunk)} chars)...")
                
                # Create extraction prompt with chunk context
                prompt = self._create_extraction_prompt(chunk, filename, chunk_index=i+1, total_chunks=len(chunks))

                logger.info("[STRATEGY_EXTRACTOR] Calling LLM for extraction...")
                response = self.llm.invoke(prompt)

                # Parse response
                if hasattr(response, 'content'):
                    response_text = response.content
                else:
                    response_text = str(response)

                logger.info(f"[STRATEGY_EXTRACTOR] LLM response (chunk {i+1}): {len(response_text)} characters")

                # Extract JSON from response
                chunk_data = self._parse_llm_response(response_text)
                logger.info(f"[STRATEGY_EXTRACTOR] Chunk {i+1}: Extracted {len(chunk_data)} records")
                all_vendor_data.extend(chunk_data)

            # Deduplicate results based on vendor_name + category + subcategory
            # This handles potential dupes from overlap
            unique_data = []
            seen_signatures = set()
            
            for record in all_vendor_data:
                # Create a simple signature
                sig = f"{record.get('vendor_name', '').lower()}|{record.get('category', '').lower()}|{record.get('subcategory', '').lower()}"
                
                # Allow duplicate vendors if strategy differs significantly, but for now simple dedup is safer
                if sig not in seen_signatures:
                    seen_signatures.add(sig)
                    unique_data.append(record)
            
            logger.info(f"[STRATEGY_EXTRACTOR] Total extracted: {len(all_vendor_data)} -> Unique: {len(unique_data)}")
            
            # Standardize keywords using StrategyKeywordStandardizer
            logger.info(f"[STRATEGY_EXTRACTOR] Standardizing {len(unique_data)} vendor records...")
            try:
                from services.strategy.keyword_standardizer import get_standardizer
                standardizer = get_standardizer()
                standardized_data = standardizer.batch_standardize(unique_data, user_id=user_id)
                logger.info(f"[STRATEGY_EXTRACTOR] ✓ Standardization complete")
                return standardized_data
            except Exception as std_error:
                logger.error(f"[STRATEGY_EXTRACTOR] Standardization failed: {std_error}")
                logger.warning("[STRATEGY_EXTRACTOR] Continuing with non-standardized data")
                return unique_data

        except Exception as e:
            logger.error(f"[STRATEGY_EXTRACTOR] LLM extraction failed: {e}")
            raise RuntimeError(f"Failed to extract vendor data: {str(e)}")

    def _create_extraction_prompt(self, document_text: str, filename: str, chunk_index: int = 1, total_chunks: int = 1) -> str:
        """Create the LLM prompt for extraction."""
        chunk_info = ""
        if total_chunks > 1:
            chunk_info = f" (Part {chunk_index} of {total_chunks}) - This is a partial extract."
            
        return f"""You are a data extraction specialist. Extract vendor procurement strategy information from the following document{chunk_info}.
            
Document: {filename}

TASK:
Extract ALL vendor entries from the document. For each vendor, extract exactly 4 fields:

1. vendor_name: The name of the vendor/manufacturer
2. category: The product category (e.g., "Flow Instruments", "Pressure Instruments", "Control Valves")
3. subcategory: The specific product subcategory (e.g., "Ultrasonic Flow Meters", "Portable Vibrometers")
4. strategy: The procurement strategy or description (e.g., "Preferred vendor for critical applications", "Sustainability and green procurement focus")

DOCUMENT CONTENT:
{document_text}

OUTPUT FORMAT:
Return ONLY a valid JSON array with NO additional text. Each object must have exactly these 4 fields:

[
  {{
    "vendor_name": "Example Vendor",
    "category": "Flow Instruments",
    "subcategory": "Ultrasonic Flow Meters",
    "strategy": "Preferred for high accuracy applications"
  }},
  ...
]

IMPORTANT:
- Extract ALL vendors found in the document
- If a field is not explicitly stated, infer it from context or use "Not specified"
- Category should be broad (e.g., "Flow Instruments", "Pressure Instruments")
- Subcategory should be specific (e.g., "Ultrasonic Flow Meters", "Differential Pressure Transmitters")
- Strategy should capture the procurement approach or vendor positioning
- Return ONLY the JSON array, no explanations or markdown formatting
"""

    def _parse_llm_response(self, response_text: str) -> List[Dict[str, str]]:
        """
        Parse LLM response to extract vendor data.

        Args:
            response_text: The LLM response text

        Returns:
            List of vendor records
        """
        # Remove markdown code blocks if present
        response_text = response_text.strip()
        if response_text.startswith('```'):
            # Remove ```json or ``` prefix
            lines = response_text.split('\n')
            response_text = '\n'.join(lines[1:-1]) if len(lines) > 2 else lines[1]

        response_text = response_text.strip()

        try:
            # Parse JSON
            data = json.loads(response_text)

            if not isinstance(data, list):
                logger.error(f"[STRATEGY_EXTRACTOR] Expected list, got {type(data)}")
                return []

            # Validate and clean each record
            cleaned_data = []
            required_fields = ['vendor_name', 'category', 'subcategory', 'strategy']

            for idx, record in enumerate(data):
                if not isinstance(record, dict):
                    logger.warning(f"[STRATEGY_EXTRACTOR] Skipping non-dict record at index {idx}")
                    continue

                # Check required fields
                missing_fields = [f for f in required_fields if f not in record]
                if missing_fields:
                    logger.warning(f"[STRATEGY_EXTRACTOR] Record {idx} missing fields: {missing_fields}")
                    # Fill in missing fields
                    for field in missing_fields:
                        record[field] = "Not specified"

                # Clean and validate
                cleaned_record = {
                    'vendor_name': str(record.get('vendor_name', 'Unknown')).strip(),
                    'category': str(record.get('category', 'Not specified')).strip(),
                    'subcategory': str(record.get('subcategory', 'Not specified')).strip(),
                    'strategy': str(record.get('strategy', 'Not specified')).strip()
                }

                # Skip empty records
                if cleaned_record['vendor_name'] in ['', 'Unknown', 'Not specified']:
                    continue

                cleaned_data.append(cleaned_record)

            logger.info(f"[STRATEGY_EXTRACTOR] Validated {len(cleaned_data)} records")
            return cleaned_data

        except json.JSONDecodeError as e:
            logger.error(f"[STRATEGY_EXTRACTOR] JSON parse error: {e}")
            logger.error(f"[STRATEGY_EXTRACTOR] Response text: {response_text[:500]}")

            # Try to extract JSON from text (sometimes LLM adds explanations)
            import re
            json_match = re.search(r'\[\s*\{.*\}\s*\]', response_text, re.DOTALL)
            if json_match:
                try:
                    data = json.loads(json_match.group(0))
                    return self._parse_llm_response(json_match.group(0))
                except:
                    pass

            return []


# Singleton instance
_extractor_instance: Optional[StrategyDocumentExtractor] = None


def get_strategy_extractor() -> StrategyDocumentExtractor:
    """
    Get singleton instance of StrategyDocumentExtractor.

    Returns:
        StrategyDocumentExtractor instance
    """
    global _extractor_instance
    if _extractor_instance is None:
        _extractor_instance = StrategyDocumentExtractor()
    return _extractor_instance
