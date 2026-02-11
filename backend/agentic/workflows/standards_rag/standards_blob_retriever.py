# agentic/workflows/standards_rag/standards_blob_retriever.py
"""
Standards Document Retriever from Azure Blob Storage

Replaces Pinecone vector store with Azure Blob-based retrieval.
Standards documents are stored as .docx files in Azure Blob, loaded and
searched using keyword matching and semantic similarity (if embeddings available).

FALLBACK STRATEGY (2026-02-11):
If Azure Blob documents don't exist (BlobNotFound), falls back to Deep Agent
local standards which have the same content but different filenames.
"""

import logging
import io
import re
from typing import List, Dict, Any, Optional
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Standards document filenames in Azure Blob 'standards-documents' container
STANDARDS_FILES = {
    "pressure": "Pressure_Instruments_Standards.docx",
    "temperature": "Temperature_Instruments_Standards.docx",
    "flow": "Flow_Instruments_Standards.docx",
    "level": "Level_Instruments_Standards.docx",
    "analytical": "Analytical_Instruments_Standards.docx",
    "safety": "Safety_Instruments_Standards.docx",
    "control_valves": "Control_Valves_Standards.docx",
    "calibration": "Calibration_Maintenance_Standards.docx",
    "communication": "Communication_Signal_Standards.docx",
    "condition_monitoring": "Condition_Monitoring_Standards.docx",
    "accessories": "Accessories_Standards.docx",
}

# Keywords for matching standards types
STANDARDS_KEYWORDS = {
    "sil": ["sil-1", "sil-2", "sil-3", "sil-4", "safety integrity level", "functional safety"],
    "iec": ["iec 61508", "iec 61511", "iec 62061", "iec"],
    "iso": ["iso 9001", "iso 14001", "iso 45001", "iso"],
    "api": ["api 670", "api 682", "api"],
    "atex": ["atex", "explosive atmosphere", "zone 0", "zone 1", "zone 2"],
    "iecex": ["iecex", "explosive", "hazardous area"],
    "asme": ["asme", "pressure vessel"],
    "ansi": ["ansi", "american national standard"],
}


class StandardsBlobRetriever:
    """Retrieves standards documents from Azure Blob Storage with Deep Agent fallback."""

    # Mapping Azure Blob filenames → Deep Agent standard types
    BLOB_TO_DEEP_AGENT_MAP = {
        "Safety_Instruments_Standards.docx": "safety",
        "Pressure_Instruments_Standards.docx": "pressure",
        "Temperature_Instruments_Standards.docx": "temperature",
        "Flow_Instruments_Standards.docx": "flow",
        "Level_Instruments_Standards.docx": "level",
        "Control_Valves_Standards.docx": "valves",
        "Analytical_Instruments_Standards.docx": "analytical",
        "Calibration_Maintenance_Standards.docx": "calibration",
        "Communication_Signal_Standards.docx": "communication",
        "Condition_Monitoring_Standards.docx": "condition_monitoring",
        "Accessories_Standards.docx": "accessories",
    }

    def __init__(self):
        """Initialize the retriever."""
        self._document_cache = {}
        self._blob_manager = None
        self._deep_agent_loader = None

    def _get_blob_manager(self):
        """Lazy load Azure Blob File Manager."""
        if self._blob_manager is None:
            from core.azure_blob_file_manager import AzureBlobFileManager
            self._blob_manager = AzureBlobFileManager()
        return self._blob_manager

    def _get_deep_agent_loader(self):
        """
        Lazy load Deep Agent document loader for fallback.

        Returns the load_standard_text function from Deep Agent.
        """
        if self._deep_agent_loader is None:
            try:
                from agentic.deep_agent.standards_deep_agent import load_standard_text
                self._deep_agent_loader = load_standard_text
                logger.info("[StandardsBlobRetriever] Deep Agent fallback available")
            except ImportError as e:
                logger.warning(f"[StandardsBlobRetriever] Deep Agent not available: {e}")
                self._deep_agent_loader = None
        return self._deep_agent_loader

    def retrieve_documents(
        self,
        question: str,
        top_k: int = 5,
        source_filter: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Retrieve relevant standards documents from Azure Blob.

        Args:
            question: User's question
            top_k: Maximum number of document chunks to return
            source_filter: Optional list of document types to search (e.g., ["safety", "pressure"])

        Returns:
            Dict with 'results' list and metadata
        """
        logger.info(f"[StandardsBlobRetriever] Retrieving documents for: {question[:100]}...")

        try:
            blob_manager = self._get_blob_manager()
            container_name = blob_manager.CONTAINERS['standards_documents']

            # Determine which documents to search
            if source_filter:
                document_types = source_filter
            else:
                # Auto-detect relevant documents from question keywords
                document_types = self._detect_relevant_documents(question)

            if not document_types:
                # Default: search all documents
                document_types = list(STANDARDS_FILES.keys())

            logger.info(f"[StandardsBlobRetriever] Searching {len(document_types)} document types: {document_types}")

            # Load and search documents
            all_results = []
            for doc_type in document_types:
                filename = STANDARDS_FILES.get(doc_type)
                if not filename:
                    continue

                # Load document content
                content = self._load_document(filename, container_name, blob_manager)
                if not content:
                    logger.warning(f"[StandardsBlobRetriever] Failed to load {filename}")
                    continue

                # Extract relevant sections
                chunks = self._extract_relevant_chunks(content, question, doc_type)
                all_results.extend(chunks)

            # Sort by relevance score and limit
            all_results.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
            top_results = all_results[:top_k]

            logger.info(f"[StandardsBlobRetriever] Found {len(top_results)} relevant chunks from {len(document_types)} documents")

            return {
                "success": True,
                "results": top_results,
                "result_count": len(top_results),
                "source": "azure_blob",
                "documents_searched": len(document_types)
            }

        except Exception as e:
            logger.error(f"[StandardsBlobRetriever] Error retrieving documents: {e}")
            return {
                "success": False,
                "results": [],
                "result_count": 0,
                "error": str(e)
            }

    def _load_document(self, filename: str, container_name: str, blob_manager) -> Optional[str]:
        """
        Load document content from Azure Blob with Deep Agent fallback.

        Strategy:
        1. Try Azure Blob first (primary source)
        2. If BlobNotFound or download fails, fall back to Deep Agent local standards
        3. Cache the result regardless of source
        """
        # Check cache first
        cache_key = f"{container_name}/{filename}"
        if cache_key in self._document_cache:
            return self._document_cache[cache_key]

        # Try Azure Blob first
        try:
            # Download from blob
            file_bytes = blob_manager.download_file(filename, container_name)
            if file_bytes:
                # Parse DOCX
                doc = DocxDocument(io.BytesIO(file_bytes))

                # Extract all text
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            paragraphs.append(" | ".join(row_text))

                full_text = "\n\n".join(paragraphs)

                # Cache it
                self._document_cache[cache_key] = full_text

                logger.debug(f"[StandardsBlobRetriever] ✓ Loaded {len(paragraphs)} paragraphs from Azure Blob: {filename}")
                return full_text

        except Exception as e:
            logger.warning(f"[StandardsBlobRetriever] Azure Blob failed for {filename}: {e}")

        # Fallback to Deep Agent local standards
        logger.info(f"[StandardsBlobRetriever] Attempting Deep Agent fallback for {filename}")

        try:
            deep_agent_loader = self._get_deep_agent_loader()
            if not deep_agent_loader:
                logger.error(f"[StandardsBlobRetriever] Deep Agent not available, cannot load {filename}")
                return None

            # Map Azure Blob filename to Deep Agent standard type
            standard_type = self.BLOB_TO_DEEP_AGENT_MAP.get(filename)
            if not standard_type:
                logger.warning(f"[StandardsBlobRetriever] No Deep Agent mapping for {filename}")
                return None

            # Load from Deep Agent
            local_content = deep_agent_loader(standard_type)
            if local_content:
                # Cache it
                self._document_cache[cache_key] = local_content
                logger.info(f"[StandardsBlobRetriever] ✓ Loaded from Deep Agent fallback: {filename} ({len(local_content)} chars)")
                return local_content
            else:
                logger.error(f"[StandardsBlobRetriever] Deep Agent also failed for {standard_type}")
                return None

        except Exception as e:
            logger.error(f"[StandardsBlobRetriever] Deep Agent fallback error for {filename}: {e}")
            return None

    def _detect_relevant_documents(self, question: str) -> List[str]:
        """Detect which document types are relevant based on question keywords."""
        question_lower = question.lower()
        relevant_types = set()

        # Check for product type keywords
        type_keywords = {
            "pressure": ["pressure", "transmitter", "gauge", "manometer"],
            "temperature": ["temperature", "thermocouple", "rtd", "thermometer"],
            "flow": ["flow", "meter", "flowmeter", "coriolis", "ultrasonic", "magnetic"],
            "level": ["level", "radar", "ultrasonic level", "guided wave"],
            "analytical": ["analytical", "analyzer", "ph", "conductivity", "gas chromatograph"],
            "safety": ["safety", "sil", "functional safety", "sis", "emergency shutdown"],
            "control_valves": ["valve", "actuator", "control valve", "ball valve", "globe valve"],
            "calibration": ["calibration", "maintenance", "testing", "verification"],
            "communication": ["communication", "protocol", "modbus", "hart", "profibus", "foundation fieldbus"],
            "condition_monitoring": ["condition monitoring", "vibration", "asset health", "predictive maintenance"],
            "accessories": ["accessory", "mounting", "bracket", "cable", "conduit"],
        }

        for doc_type, keywords in type_keywords.items():
            if any(kw in question_lower for kw in keywords):
                relevant_types.add(doc_type)

        # Check for standards keywords (always include safety if standards mentioned)
        for standard_type, keywords in STANDARDS_KEYWORDS.items():
            if any(kw in question_lower for kw in keywords):
                relevant_types.add("safety")  # Safety document usually has standards info
                break

        return list(relevant_types) if relevant_types else []

    def _extract_relevant_chunks(self, content: str, question: str, doc_type: str) -> List[Dict[str, Any]]:
        """Extract relevant text chunks from document content."""
        # Split into paragraphs
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

        # Keywords from question
        question_lower = question.lower()
        question_words = set(re.findall(r'\b\w+\b', question_lower))

        chunks = []
        for i, para in enumerate(paragraphs):
            para_lower = para.lower()

            # Calculate simple relevance score
            score = 0.0

            # Check keyword overlap
            para_words = set(re.findall(r'\b\w+\b', para_lower))
            overlap = len(question_words & para_words)
            score += overlap * 0.1

            # Boost for standards mentions
            for standard_keywords in STANDARDS_KEYWORDS.values():
                if any(kw in para_lower for kw in standard_keywords):
                    score += 0.5

            # Boost for exact question terms
            if any(word in para_lower for word in question_words if len(word) > 3):
                score += 0.3

            # Only include chunks with some relevance
            if score > 0.2 or len(para) > 200:  # Include long paragraphs even with low score
                chunks.append({
                    "id": f"{doc_type}_chunk_{i}",
                    "content": para,
                    "metadata": {
                        "filename": STANDARDS_FILES.get(doc_type, "unknown"),
                        "standard_type": doc_type,
                        "chunk_index": i
                    },
                    "relevance_score": score
                })

        return chunks


# Singleton instance
_retriever_instance = None


def get_standards_blob_retriever() -> StandardsBlobRetriever:
    """Get singleton instance of StandardsBlobRetriever."""
    global _retriever_instance
    if _retriever_instance is None:
        _retriever_instance = StandardsBlobRetriever()
    return _retriever_instance
